"""Independent Report Loop orchestration service."""
from __future__ import annotations

import copy
import difflib
import hashlib
import json
import shutil
import threading
import time
import uuid
from pathlib import Path
from typing import Callable, Dict, Optional

HERE = Path(__file__).resolve().parent
ROOT = HERE.parent
HARNESS = ROOT / "harness"
import sys
if str(HARNESS) not in sys.path:
    sys.path.insert(0, str(HARNESS))

import llm_client
from report_failure import failure_report_from_checks
from report_scoring import normalize_check_scores, score_labeled_check_judgment
from data_packages import load_data_package_case, resolve_data_json
from report_loop_settings import ReportLoopSettings
from judge_batch import (
    DEFAULT_JUDGE_MAX_RETRIES,
    JUDGE_STRATEGY_PER_DIMENSION,
    judge_report,
)
from judge_prompt import build_judge_prompt
from report_loop_state import (
    actions,
    enforce_pre_generation_budget,

    next_version,
    public_run,

    revision,

    settle_judged_revision,
)
from report_loop_store import ReportLoopStore
from report_loop_utils import directory_hash
from external_run_models import ExternalRunRequest, ReportOutputContract
from report_source import report_text_from_result
from workbuddy_runner import run_external_cases
from workbuddy_batch.dataset import load_openharness_rows
from workbuddy_batch.io import sha256_text


_MAX_REJECTED_ATTEMPTS = 2
_MAX_AVOID_PATTERNS_PER_CHECK = 2
_MAX_REJECTED_DIFF_CHARS = 4000


class ReportLoopError(RuntimeError):
    pass


def _sha_file(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _now() -> float:
    return round(time.time(), 3)


class ReportLoopService:
    """Owns Report Loop state; never reads or writes Skill Loop sessions."""

    def __init__(
        self,
        root: Path,
        data_root: Path,
        skills_root: Path,
        rubric_path: Path,
        settings: Optional[ReportLoopSettings] = None,
        runner_func: Optional[Callable] = None,
        judge_func: Optional[Callable] = None,
        call_llm_func: Optional[Callable] = None,
    ):
        self.store = ReportLoopStore(root)
        self.data_root = Path(data_root).resolve()
        self.skills_root = Path(skills_root).resolve()
        self.rubric_path = Path(rubric_path).resolve()
        self.settings = settings or ReportLoopSettings.from_env()
        self.runner_func = runner_func or run_external_cases
        self.judge_func = judge_func or judge_report
        self.call_llm_func = call_llm_func or llm_client.call_llm
        self._lock = threading.RLock()
        self._runs = {item["id"]: item for item in self.store.load_all()}
        self._recover_interrupted_jobs()
        self._resume_token_budget_stops()

    def _resume_token_budget_stops(self) -> None:
        """Reopen runs stopped by the retired Token-budget policy."""
        for run in self._runs.values():
            stop_state = run.get("stop_state") or {}
            if stop_state.get("code") != "token_budget_exhausted":
                continue
            run["stop_state"] = {
                "stopped": False,
                "code": None,
                "reason": None,
            }
            run["status"] = self._stable_status(run)
            self.store.append_event(run["id"], "run.resumed", {
                "reason": "token_budget_limit_removed",
            })
            self.store.save(run)

    def _recover_interrupted_jobs(self) -> None:
        for run in self._runs.values():
            changed = False
            for job in (run.get("jobs") or {}).values():
                if job.get("status") not in {"queued", "running"}:
                    continue
                recovered = (
                    job.get("kind") == "generate"
                    and self._recover_generated_job(run, job)
                )
                job.update({
                    "status": "completed" if recovered else "failed",
                    "finished_at": _now(),
                    "error": None if recovered else "服务重启，后台任务已中断",
                })
                changed = True
            if changed:
                run["status"] = self._stable_status(run)
                self.store.save(run)

    def _recover_generated_job(self, run: Dict, job: Dict) -> bool:
        """Import a fully captured Runner artifact left by an interrupted job."""
        version = str(job.get("target_version") or "")
        if not version or next_version(run) != version:
            return False
        run_directory = self.store.run_dir(run["id"])
        runner_root = run_directory / "runner"
        if not runner_root.is_dir():
            return False
        started_at = float(job.get("started_at") or job.get("created_at") or 0)
        candidates = sorted(
            runner_root.glob("gen-*"),
            key=lambda path: path.stat().st_mtime,
            reverse=True,
        )
        for generation_dir in candidates:
            try:
                request = json.loads(
                    (generation_dir / "request.json").read_text(encoding="utf-8")
                )
            except (OSError, UnicodeError, ValueError):
                continue
            if (
                request.get("session_id") != run["id"]
                or request.get("skill_version") != "report-" + version
                or run["case_id"] not in (request.get("openharness_case_ids") or [])
                or generation_dir.stat().st_mtime < started_at - 5
            ):
                continue
            for attempt in sorted(generation_dir.glob("attempt*"), reverse=True):
                report_path = attempt / "artifacts" / "report.md"
                manifest_path = attempt / "artifacts" / "manifest.json"
                try:
                    raw_report = report_path.read_text(encoding="utf-8-sig")
                    report_text = raw_report.strip()
                    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
                except (OSError, UnicodeError, ValueError):
                    continue
                report_entry = next(
                    (
                        item for item in manifest
                        if item.get("path") == "deliverables/report.md"
                    ),
                    None,
                )
                if (
                    not (attempt / "conversation.md").is_file()
                    or not report_entry
                    or report_entry.get("status") == "deleted"
                    or report_entry.get("sha256") != sha256_text(raw_report)
                    or len(report_text.encode("utf-8")) < self.settings.min_report_bytes
                ):
                    continue
                parent_version = None if version == "v1" else run.get("current_version")
                run["revisions"].append({
                    "version": version,
                    "parent_version": parent_version,
                    "status": "report_ready",
                    "decision": "pending",
                    "score_delta": None,
                    "report_sha256": sha256_text(report_text),
                    "report_path": report_path.relative_to(run_directory).as_posix(),
                    "generation_job_id": job["id"],
                    "generation_id": request.get("generation_id") or generation_dir.name,
                    "runner_status": "generated",
                    "trace": {
                        "output_dir": generation_dir.relative_to(run_directory).as_posix(),
                        "case": {"status": "generated", "recovered": True},
                    },
                    "judgment": None,
                    "failure_report": None,
                    "created_at": _now(),
                })
                run["status"] = "report_ready"
                self.store.append_event(run["id"], "report.generated", {
                    "version": version,
                    "parent_version": parent_version,
                    "backend": "workbuddy",
                    "report_sha256": sha256_text(report_text),
                    "recovered": True,
                })
                return True
        return False

    @staticmethod
    def _stable_status(run: Dict) -> str:
        if (run.get("stop_state") or {}).get("stopped"):
            return "completed"
        revisions = run.get("revisions") or []
        if not revisions:
            return "imported"
        if revisions[-1].get("status") == "report_ready":
            return "report_ready"
        return "judged"

    def _get(self, run_id: str) -> Dict:
        run = self._runs.get(str(run_id or ""))
        if run is None:
            raise ReportLoopError("Report Run 不存在: %s" % run_id)
        return run

    def _public(self, run: Dict) -> Dict:
        return public_run(
            run,
            lambda version: self.store.load_report(run["id"], version),
        )

    def view(self, run_id: str) -> Dict:
        with self._lock:
            return self._public(self._get(run_id))

    def list_runs(self) -> list[Dict]:
        with self._lock:
            return [
                self._public(item)
                for item in sorted(
                    self._runs.values(),
                    key=lambda value: float(value.get("created_at") or 0),
                    reverse=True,
                )
            ]

    def generation_chain(self, run_id: str, version: str) -> Dict:
        """Locate the runner ``conversation.md`` for a revision and parse it into a
        3-layer execution chain (conversation / execution-chain / execution-detail)."""
        run = self.view(run_id)
        rev = next((r for r in (run.get("revisions") or []) if r.get("version") == version), None)
        if rev is None:
            raise ReportLoopError("版本不存在: " + str(version))
        run_dir = self.store.run_dir(run_id)
        text, rel = self._locate_conversation(run_dir, rev)
        if not text:
            return {
                "found": False,
                "note": "该版本未生成本地运行轨迹（可能经 API 生成，无 runner 目录），以下为原始 trace。",
                "fallback_trace": rev.get("trace"),
            }
        return {"found": True, "path": rel, "chain": self.parse_conversation(text)}

    def _locate_conversation(self, run_dir: "Path", rev: Dict):
        runner = run_dir / "runner"
        candidates = list(runner.rglob("conversation.md")) if runner.exists() else []
        output_dir = (rev.get("trace") or {}).get("output_dir")
        gen_id = rev.get("generation_id")
        chosen = None
        if output_dir:
            for cand in (run_dir / output_dir / "attempt1" / "conversation.md",
                         run_dir / output_dir / "conversation.md"):
                if cand.exists():
                    chosen = cand
                    break
        if chosen is None and gen_id:
            for cand in (run_dir / "runner" / gen_id / "attempt1" / "conversation.md",
                         run_dir / "runner" / gen_id / "conversation.md"):
                if cand.exists():
                    chosen = cand
                    break
        if chosen is None and len(candidates) == 1:
            chosen = candidates[0]
        if chosen is None and gen_id and candidates:
            for cand in candidates:
                if gen_id in str(cand):
                    chosen = cand
                    break
        if chosen is None and candidates:
            chosen = candidates[0]
        if chosen is None:
            return None, None
        rel = str(chosen.relative_to(run_dir))
        try:
            return chosen.read_text(encoding="utf-8"), rel
        except OSError:
            return None, None

    @staticmethod
    def _strip_fence(block: str) -> str:
        lines = block.split("\n")
        while lines and not lines[0].strip():
            lines.pop(0)
        while lines and not lines[-1].strip():
            lines.pop()
        if lines and lines[0].lstrip().startswith("```") and lines[-1].lstrip() == "```":
            lines = lines[1:-1]
        return "\n".join(lines).strip("\n")

    @staticmethod
    def parse_conversation(text: str) -> Dict:
        import re  # noqa
        lines = text.replace("\r\n", "\n").split("\n")
        turns: list = []
        cur_turn = None
        cur_tool = None
        tool_section = None   # None | "meta" | "input" | "output"
        tool_prefix = ""      # "> " to strip in old quoted format
        tool_buf: list = []
        in_thinking = False
        think_buf: list = []

        def flush_tool():
            if cur_tool is None:
                return
            if tool_section in ("input", "output") and tool_buf:
                content = ReportLoopService._strip_fence("\n".join(tool_buf).strip("\n"))
                if tool_section == "input":
                    cur_tool["input"] = content
                else:
                    cur_tool["output"] = content
            tool_buf.clear()

        def end_tool():
            flush_tool()
            nonlocal cur_tool, tool_section, tool_prefix
            cur_tool = None
            tool_section = None
            tool_prefix = ""

        def end_thinking():
            nonlocal in_thinking, think_buf
            if in_thinking and cur_turn is not None:
                cur_turn["thinking"] = (cur_turn.get("thinking") or "") + "\n".join(think_buf).strip("\n")
            in_thinking = False
            think_buf = []

        for raw in lines:
            stripped = raw.strip()
            if not stripped:
                if in_thinking:
                    think_buf.append("")
                elif tool_section in ("input", "output"):
                    tool_buf.append("")
                continue

            if stripped.startswith("**user:**"):
                end_thinking(); end_tool()
                cur_turn = {"role": "user", "content": stripped[len("**user:**"):].strip()}
                turns.append(cur_turn)
                continue
            if stripped.startswith("**workbuddy:**"):
                end_thinking(); end_tool()
                if not cur_turn or cur_turn.get("role") != "assistant":
                    cur_turn = {"role": "assistant", "content": "", "thinking": None, "tools": []}
                    turns.append(cur_turn)
                cur_turn["content"] = (cur_turn["content"] or "") + stripped[len("**workbuddy:**"):].strip() + "\n"
                continue
            if stripped.startswith("**tool:"):
                end_thinking(); end_tool()
                if not cur_turn or cur_turn.get("role") != "assistant":
                    cur_turn = {"role": "assistant", "content": "", "thinking": None, "tools": []}
                    turns.append(cur_turn)
                m = re.match(r"\*\*tool:([A-Za-z_][\w.]*)\((.*)\)\*\*", stripped)
                name = m.group(1) if m else "tool"
                args = m.group(2) if m else ""
                cur_tool = {"name": name, "args": args, "status": None, "duration": None,
                            "input": None, "output": None}
                cur_turn["tools"].append(cur_tool)
                tool_section = "meta"
                tool_buf = []
                continue
            if stripped.startswith("**thinking:**") or stripped.startswith("**思考:**") or stripped.startswith("**思考**"):
                end_thinking(); end_tool()
                if not cur_turn or cur_turn.get("role") != "assistant":
                    cur_turn = {"role": "assistant", "content": "", "thinking": None, "tools": []}
                    turns.append(cur_turn)
                if cur_turn.get("thinking") is None:
                    cur_turn["thinking"] = ""
                in_thinking = True
                think_buf = []
                continue
            # old quoted input/output markers ( "> **Input**" / "> **Result**" )
            if stripped.startswith("> **Input**") or stripped.startswith("> **输入**"):
                end_thinking()
                if cur_tool is not None:
                    flush_tool(); tool_section = "input"; tool_prefix = "> "; tool_buf = []
                continue
            if stripped.startswith("> **Result**") or stripped.startswith("> **结果**") or stripped.startswith("> **Output**"):
                end_thinking()
                if cur_tool is not None:
                    flush_tool(); tool_section = "output"; tool_prefix = "> "; tool_buf = []
                continue
            # new input/output markers ( "input:" / "result:" )
            if re.match(r"^(input|result|output)\s*:\s*$", stripped, re.IGNORECASE):
                end_thinking()
                if cur_tool is not None:
                    flush_tool()
                    sec = stripped.split(":")[0].lower()
                    tool_section = "output" if sec == "result" else sec
                    tool_prefix = ""
                    tool_buf = []
                continue

            if in_thinking:
                think_buf.append(re.sub(r"^>\s?", "", raw))
                continue
            if cur_tool is not None:
                if tool_section == "meta":
                    sm = re.search(r"`(success|error|failed|timeout)`", stripped, re.IGNORECASE)
                    dm = re.search(r"duration:\s*`?([^`\n]+?)`?\s*$", stripped)
                    if sm:
                        cur_tool["status"] = sm.group(1).lower()
                    if dm:
                        cur_tool["duration"] = dm.group(1).strip()
                    continue
                if tool_section in ("input", "output"):
                    if tool_prefix and raw.lstrip().startswith(">"):
                        tool_buf.append(re.sub(r"^\s*>\s?", "", raw))
                    else:
                        tool_buf.append(raw)
                    continue
            if cur_turn is not None:
                cur_turn["content"] = (cur_turn["content"] or "") + raw + "\n"
        end_thinking(); end_tool()
        for t in turns:
            if t.get("role") == "assistant":
                t["content"] = (t.get("content") or "").strip("\n")
                if not t.get("tools"):
                    t.pop("tools", None)
                if t.get("thinking") in (None, ""):
                    t.pop("thinking", None)
        return {"turns": turns}

    def job(self, job_id: str) -> Dict:
        with self._lock:
            for run in self._runs.values():
                job = (run.get("jobs") or {}).get(job_id)
                if job:
                    return copy.deepcopy(job)
        raise ReportLoopError("Report Loop 任务不存在: %s" % job_id)

    def create_run(
        self,
        data_id: str,
        case_id: str,
        skill_template_id: str,
        requirement: str = "",
        creator: str = "local",
        overall_target: float = 5.0,
        max_no_improvement: int = 2,
        max_elapsed_seconds: int = 3600,
        stop_on_unrepairable_failure: bool = False,
    ) -> Dict:
        target = float(overall_target)
        patience = int(max_no_improvement)
        time_budget = int(max_elapsed_seconds)
        if not 0 <= target <= 5:
            raise ValueError("loop 停止总分必须在 0 到 5 之间")
        if patience < 1:
            raise ValueError("连续无提升版本数必须至少为 1")
        if time_budget < 1:
            raise ValueError("Time budget 必须为正整数秒")
        dataset_path = resolve_data_json(self.data_root, data_id)
        rows = load_openharness_rows(dataset_path)
        case = next(
            (row for row in rows if str(row.get("case_id")) == str(case_id)),
            None,
        )
        if case is None:
            raise ValueError("数据集不包含 Case: %s" % case_id)
        preview = load_data_package_case(self.data_root, data_id, case_id)
        skill_id = str(skill_template_id or "").strip()
        if not skill_id or Path(skill_id).name != skill_id:
            raise ValueError("非法 Skill 模板: %s" % skill_id)
        skill_source = (self.skills_root / skill_id).resolve()
        skill_source.relative_to(self.skills_root)
        if not (skill_source / "SKILL.md").is_file():
            raise ValueError("Skill 模板缺少 SKILL.md: %s" % skill_id)
        if not self.rubric_path.is_file():
            raise ValueError("Report Loop rubric 不存在: %s" % self.rubric_path)
        rubric = json.loads(self.rubric_path.read_text(encoding="utf-8"))
        run_id = "report-" + uuid.uuid4().hex[:8]
        created_at = _now()
        frozen_skill_path = self.store.run_dir(run_id) / "skill"
        run = {
            "id": run_id,
            "loop_kind": "report",
            "creator": str(creator or "local"),
            "requirement": str(requirement or "").strip(),
            "data_id": data_id,
            "data_path": str(dataset_path),
            "data_sha256": _sha_file(dataset_path),
            "case_id": str(case_id),
            "topic": preview.get("topic") or str(case_id),
            "case": case,
            "dataset_preview": preview,
            "skill_template_id": skill_id,
            "frozen_skill_path": str(frozen_skill_path),
            "frozen_skill_sha256": directory_hash(skill_source),
            "rubric": rubric,
            "rubric_path": str(self.rubric_path),
            "rubric_sha256": _sha_file(self.rubric_path),
            "stop_policy": {
                "overall_target": target,
                "max_no_improvement": patience,
                "max_elapsed_seconds": time_budget,
                "stop_on_unrepairable_failure": bool(stop_on_unrepairable_failure),
            },
            "stop_state": {"stopped": False, "reason": None},
            "no_improvement_streak": 0,
            "current_version": None,
            "status": "imported",
            "revisions": [],
            "jobs": {},
            "created_at": created_at,
            "updated_at": created_at,
        }
        with self._lock:
            self.store.initialize(run)
            try:
                shutil.copytree(skill_source, frozen_skill_path)
                run["frozen_skill_sha256"] = directory_hash(frozen_skill_path)
                self._runs[run_id] = run
                self.store.save(run)
                self.store.append_event(run_id, "run.created", {
                    "data_id": data_id,
                    "case_id": case_id,
                    "skill_template_id": skill_id,
                })
            except Exception:
                shutil.rmtree(self.store.run_dir(run_id), ignore_errors=True)
                raise
            return self._public(run)

    def _start_job(self, run_id: str, kind: str, target_version: str, worker) -> Dict:
        with self._lock:
            run = self._get(run_id)
            if any(
                item.get("status") in {"queued", "running"}
                for item in run.get("jobs", {}).values()
            ):
                raise ReportLoopError("当前 Report Run 已有执行中的任务")
            job_id = "%s-%s" % (kind, uuid.uuid4().hex[:10])
            job = {
                "id": job_id,
                "run_id": run_id,
                "kind": kind,
                "target_version": target_version,
                "status": "queued",
                "created_at": _now(),
                "started_at": None,
                "finished_at": None,
                "error": None,
            }
            run.setdefault("jobs", {})[job_id] = job
            run["status"] = kind + "_queued"
            self.store.save(run)
        thread = threading.Thread(
            target=self._run_job,
            args=(run_id, job_id, worker),
            daemon=True,
            name="report-loop-" + job_id,
        )
        thread.start()
        return copy.deepcopy(job)

    def _run_job(self, run_id: str, job_id: str, worker) -> None:
        with self._lock:
            run = self._get(run_id)
            job = run["jobs"][job_id]
            job.update({"status": "running", "started_at": _now()})
            run["status"] = job["kind"] + "_running"
            self.store.save(run)
        try:
            worker(run_id, job_id)
            with self._lock:
                run = self._get(run_id)
                job = run["jobs"][job_id]
                job.update({"status": "completed", "finished_at": _now()})
                self.store.save(run)
        except Exception as exc:
            with self._lock:
                run = self._get(run_id)
                job = run["jobs"][job_id]
                job.update({
                    "status": "failed",
                    "finished_at": _now(),
                    "error": str(exc),
                })
                run["status"] = self._stable_status(run)
                self.store.append_event(run_id, "job.failed", {
                    "job_id": job_id,
                    "kind": job["kind"],
                    "error": str(exc),
                })
                self.store.save(run)

    def _runner_model(self, model: Optional[str]) -> str:
        selected = str(model or self.settings.model or "").strip()
        if not selected:
            raise ReportLoopError("Runner model is required")
        if selected not in self.settings.models:
            raise ReportLoopError(
                "Unsupported WorkBuddy Runner model: %s" % selected
            )
        return selected

    def start_generate(
        self,
        run_id: str,
        model: Optional[str] = None,
        generation_backend: Optional[str] = None,
        api_model: Optional[str] = None,
        reasoning_effort: Optional[str] = None,
    ) -> Dict:
        backend = str(generation_backend or "workbuddy").strip().lower()
        if backend not in {"workbuddy", "api", "codex"}:
            raise ReportLoopError("不支持的报告生成调用方式: %s" % backend)
        if backend == "workbuddy":
            model = self._runner_model(model)
        else:
            # bianxie API / codex: 直接调用 LLM 编写报告
            chosen = str(model or api_model or "").strip()
            if not chosen:
                raise ReportLoopError("bianxie API 报告生成需要指定 API 模型")
            model = chosen
        with self._lock:
            run = self._get(run_id)
            if enforce_pre_generation_budget(run):
                self.store.save(run)
            action = actions(run)["generate"]
            if not action["enabled"]:
                raise ReportLoopError("当前状态不能生成下一版报告")
            parent = action.get("base_version")
            version = action["next_version"]
        return self._start_job(
            run_id,
            "generate",
            version,
            lambda rid, jid: self._generate(
                rid, jid, parent, backend, model, reasoning_effort
            ),
        )

    def start_optimize(self, run_id: str, model: Optional[str] = None) -> Dict:
        """Compatibility alias; report iteration is now a Runner generation."""
        return self.start_generate(run_id, model)

    def start_judge(
        self,
        run_id: str,
        backend: Optional[str] = None,
        model: Optional[str] = None,
        reasoning_effort: Optional[str] = None,
    ) -> Dict:
        with self._lock:
            run = self._get(run_id)
            action = actions(run)["judge"]
            if not action["enabled"]:
                raise ReportLoopError("当前状态没有待评测报告")
            version = action["version"]
        return self._start_job(
            run_id,
            "judge",
            version,
            lambda rid, jid: self._judge(
                rid, jid, version, backend, model, reasoning_effort
            ),
        )

    def _build_evidence(self, run: Dict) -> str:
        """从结构化证据索引抽取权威数值/时间范围/来源类型，供改写时对照。"""
        preview = run.get("dataset_preview") or {}
        sd = preview.get("structured_data") or {}
        items = sd.get("items") or []
        if not items:
            return ""
        lines = []
        for it in items:
            sid = it.get("id", "")
            sref = it.get("source_ref", "")
            content = it.get("content", "")
            lines.append(f"[{sid} | {sref}] {content}")
        return "\n".join(lines)

    @staticmethod
    def _check_definitions(run: Dict) -> Dict:
        definitions = {}
        for dimension in (run.get("rubric") or {}).get("dimensions", []):
            for check in dimension.get("checks") or []:
                check_id = str(check.get("id") or "").strip()
                if not check_id:
                    continue
                definitions[check_id] = {
                    "dimension": dimension.get("name"),
                    "dimension_zh": dimension.get(
                        "name_zh", dimension.get("name")
                    ),
                    "label": check.get("label", check_id),
                    "requirement": check.get("desc") or "",
                    "redline": bool(check.get("redline")),
                }
        return definitions

    @staticmethod
    def _failure_checks(failure_report: Dict) -> Dict:
        checks = (failure_report or {}).get("checks")
        if isinstance(checks, dict):
            return {
                str(check_id): value
                for check_id, value in checks.items()
                if isinstance(value, dict)
            }
        result = {}
        for dimension in (failure_report or {}).get("dimensions") or []:
            for failure in dimension.get("failures") or []:
                check_id = str(failure.get("check_id") or "").strip()
                if check_id:
                    result[check_id] = failure
        return result

    @staticmethod
    def _is_met(value) -> bool:
        return value in {"met", 1, 1.0}

    def _build_revision_plan(
        self,
        run: Dict,
        accepted: Dict,
        rejected: list,
    ) -> Dict:
        """Build disjoint current constraints plus bounded regression memory."""
        definitions = self._check_definitions(run)
        judgment = accepted.get("judgment") or {}
        current_checks = judgment.get("checks") or {}
        reasoning = judgment.get("reasoning") or {}
        current_failures = self._failure_checks(
            accepted.get("failure_report") or {}
        )
        repair_targets = {}
        preserve_constraints = {}

        for check_id, value in current_checks.items():
            check_id = str(check_id)
            definition = copy.deepcopy(definitions.get(check_id) or {})
            reason = str(reasoning.get(check_id) or "")
            if self._is_met(value):
                preserve_constraints[check_id] = {
                    "status": "met",
                    **definition,
                    "baseline_reason": reason,
                }
                continue
            failure = current_failures.get(check_id) or {}
            repair_targets[check_id] = {
                "status": failure.get("status", value),
                **definition,
                "current_problem": str(failure.get("reason") or reason),
                "historical_patterns": [],
            }

        avoid_patterns = []
        pattern_counts = {}
        seen_patterns = set()
        for rejected_item in reversed(rejected):
            failures = self._failure_checks(
                rejected_item.get("failures")
                or rejected_item.get("failure_report")
                or {}
            )
            for check_id, failure in failures.items():
                if check_id not in current_checks:
                    continue
                pattern = str(failure.get("reason") or "").strip()
                if not pattern:
                    continue
                dedupe_key = (check_id, pattern)
                if dedupe_key in seen_patterns:
                    continue
                if (
                    pattern_counts.get(check_id, 0)
                    >= _MAX_AVOID_PATTERNS_PER_CHECK
                ):
                    continue
                seen_patterns.add(dedupe_key)
                pattern_counts[check_id] = pattern_counts.get(check_id, 0) + 1
                historical = {
                    "source_version": rejected_item.get("version"),
                    "pattern": pattern,
                }
                if self._is_met(current_checks.get(check_id)):
                    avoid_patterns.append({
                        "check_id": check_id,
                        **definitions.get(check_id, {}),
                        **historical,
                        "instruction": (
                            "保持当前已采纳版本的正确做法，不得重新引入该历史回退。"
                        ),
                    })
                elif check_id in repair_targets:
                    repair_targets[check_id]["historical_patterns"].append(
                        historical
                    )

        rejected_attempts = [
            {
                key: copy.deepcopy(item.get(key))
                for key in (
                    "version",
                    "parent_version",
                    "adoption_gate",
                    "dimension_delta",
                    "rejected_diff",
                )
            }
            for item in rejected[-_MAX_REJECTED_ATTEMPTS:]
        ]
        overlap = set(repair_targets) & set(preserve_constraints)
        if overlap:
            raise ReportLoopError(
                "Rewrite 约束冲突，check 同时要求修复与保持: "
                + ", ".join(sorted(overlap))
            )
        return {
            "repair_targets": repair_targets,
            "preserve_constraints": preserve_constraints,
            "avoid_patterns": list(reversed(avoid_patterns)),
            "rejected_attempts": rejected_attempts,
        }

    @staticmethod
    def _dimension_delta(candidate: Dict, baseline: Dict) -> Dict:
        candidate_judgment = candidate.get("judgment") or {}
        baseline_judgment = baseline.get("judgment") or {}
        candidate_dimensions = candidate_judgment.get("dimensions") or {}
        baseline_dimensions = baseline_judgment.get("dimensions") or {}
        delta = {}
        for name in sorted(set(candidate_dimensions) | set(baseline_dimensions)):
            try:
                delta[name] = round(
                    float(candidate_dimensions.get(name) or 0)
                    - float(baseline_dimensions.get(name) or 0),
                    3,
                )
            except (TypeError, ValueError):
                continue
        try:
            delta["overall"] = round(
                float(candidate_judgment.get("overall") or 0)
                - float(baseline_judgment.get("overall") or 0),
                3,
            )
        except (TypeError, ValueError):
            pass
        return delta

    @staticmethod
    def _report_diff(
        baseline_text: str,
        candidate_text: str,
        baseline_version: str,
        candidate_version: str,
        max_chars: int = 12000,
    ) -> str:
        diff = "\n".join(difflib.unified_diff(
            str(baseline_text or "").splitlines(),
            str(candidate_text or "").splitlines(),
            fromfile=baseline_version,
            tofile=candidate_version,
            lineterm="",
        ))
        if len(diff) <= max_chars:
            return diff
        return diff[:max_chars].rstrip() + "\n... [diff truncated]"

    def _revision_dataset(self, run: Dict, parent_version: str, version: str) -> Path:
        case = copy.deepcopy(run["case"])
        package_root = Path(run["data_path"]).parent.resolve()
        for spec in case.get("input_files") or []:
            if not isinstance(spec, dict) or not spec.get("source"):
                continue
            source = Path(str(spec["source"]))
            if not source.is_absolute():
                spec["source"] = str((package_root / source).resolve())
        accepted = revision(run, parent_version)
        base_report = self.store.load_report(run["id"], parent_version)
        rejected = []
        for item in run.get("revisions", []):
            if item.get("decision") != "rejected":
                continue
            rejected_parent_version = item.get("parent_version")
            if not rejected_parent_version:
                continue
            rejected_parent = revision(run, rejected_parent_version)
            rejected.append({
                "version": item["version"],
                "parent_version": rejected_parent_version,
                "overall": (item.get("judgment") or {}).get("overall"),
                "failures": item.get("failure_report"),
                "adoption_gate": copy.deepcopy(item.get("adoption_gate") or {}),
                "dimension_delta": self._dimension_delta(item, rejected_parent),
                "rejected_diff": self._report_diff(
                    self.store.load_report(run["id"], rejected_parent_version),
                    self.store.load_report(run["id"], item["version"]),
                    rejected_parent_version,
                    item["version"],
                    max_chars=_MAX_REJECTED_DIFF_CHARS,
                ),
            })
        evidence = self._build_evidence(run)
        revision_plan = self._build_revision_plan(run, accepted, rejected)
        task = (
            "以当前已采纳版本为唯一事实基线，进行最小范围修订。"
            "优先修复revision_plan.repair_targets；"
            "不得删除、弱化或改写revision_plan.preserve_constraints中的已通过质量属性；"
            "revision_plan.avoid_patterns和revision_plan.rejected_attempts仅用于防回退，"
            "不得据此推翻当前判定。若历史信息与当前判定冲突，以当前判定为准。"
            "不要改写与修复目标无关的段落，不要输出修改说明，只交付完整新版报告。"
        )
        original_turns = copy.deepcopy(case.get("turns") or [])
        prompt = {
            "task": task,
            "base_version": parent_version,
            "base_report": base_report,
            "evidence": evidence,
            "revision_plan": revision_plan,
            "original_turns": original_turns,
            "requirement": run.get("requirement") or "",
        }
        case["turns"] = [{
            "round": 0,
            "label": "report_loop_runner_revision",
            "prompt": json.dumps(prompt, ensure_ascii=False),
        }]
        directory = self.store.run_dir(run["id"]) / "revision_inputs" / version
        directory.mkdir(parents=True, exist_ok=True)
        path = directory / "data.json"
        document = {
            "schema_version": "openharness-wb/v1",
            "defaults": {"skills": [run["skill_template_id"]]},
            "cases": [case],
        }
        path.write_text(json.dumps(document, ensure_ascii=False, indent=2), encoding="utf-8")
        return path

    def _report_path(
        self,
        run_id: str,
        version: str,
        case_result: Dict,
        report_text: str,
    ) -> str:
        directory = self.store.run_dir(run_id)
        captured = str(
            ((case_result.get("report") or {}).get("captured_path") or "")
        ).strip()
        if captured:
            try:
                path = Path(captured).resolve()
                relative = path.relative_to(directory)
                if (
                    path.is_file()
                    and sha256_text(path.read_text(encoding="utf-8-sig"))
                    == sha256_text(report_text)
                ):
                    return relative.as_posix()
            except (OSError, UnicodeError, ValueError):
                pass

        # Custom/test runners may return text without a captured artifact.
        return self.store.save_report(
            run_id, version, report_text
        ).relative_to(directory).as_posix()

    def _compact_runner_trace(self, value, directory: Path):
        if isinstance(value, list):
            return [
                self._compact_runner_trace(item, directory) for item in value
            ]
        if not isinstance(value, dict):
            return value
        compact = {}
        for key, item in value.items():
            if key == "text" and "captured_path" in value:
                continue
            if key in {
                "captured_path", "run_dir", "trace_path", "manifest_path"
            }:
                try:
                    item = (
                        Path(str(item))
                        .resolve()
                        .relative_to(directory)
                        .as_posix()
                    )
                except (OSError, ValueError):
                    pass
            compact[key] = self._compact_runner_trace(item, directory)
        return compact

    def _generate(
        self,
        run_id: str,
        job_id: str,
        parent_version: Optional[str],
        generation_backend: str = "workbuddy",
        model: Optional[str] = None,
        reasoning_effort: Optional[str] = None,
    ) -> None:
        with self._lock:
            run = copy.deepcopy(self._get(run_id))
            version = self._get(run_id)["jobs"][job_id]["target_version"]
        case_file = (
            Path(run["data_path"])
            if parent_version is None
            else self._revision_dataset(run, parent_version, version)
        )
        if generation_backend == "workbuddy":
            report_text, trace, generation_id = self._generate_via_workbuddy(
                run_id, run, version, case_file, model
            )
        else:
            report_text, trace, generation_id = self._generate_via_api(
                run, version, parent_version, case_file, generation_backend, model, reasoning_effort
            )
        if len(report_text.encode("utf-8")) < self.settings.min_report_bytes:
            raise ReportLoopError("报告长度不足，生成失败")
        run_directory = self.store.run_dir(run_id)
        if generation_backend == "workbuddy":
            report_path = self._report_path(
                run_id, version, trace.get("case") or {}, report_text
            )
        else:
            report_path = (
                self.store.save_report(run_id, version, report_text)
                .relative_to(run_directory)
                .as_posix()
            )
        item = {
            "version": version,
            "parent_version": parent_version,
            "status": "report_ready",
            "decision": "pending",
            "score_delta": None,
            "report_sha256": sha256_text(report_text),
            "report_path": report_path,
            "generation_job_id": job_id,
            "generation_id": generation_id,
            "runner_status": "generated",
            "trace": trace,
            "judgment": None,
            "failure_report": None,
            "created_at": _now(),
        }
        with self._lock:
            current = self._get(run_id)
            if next_version(current) != version:
                raise ReportLoopError("报告版本已被并发修改")
            current["revisions"].append(item)
            current["status"] = "report_ready"
            self.store.append_event(run_id, "report.generated", {
                "version": version,
                "parent_version": parent_version,
                "backend": generation_backend,
                "report_sha256": item["report_sha256"],
            })
            self.store.save(current)

    def _generate_via_workbuddy(self, run_id, run, version, case_file, model):
        run_directory = self.store.run_dir(run_id)
        request = ExternalRunRequest(
            case_file=case_file,
            output_root=run_directory / "runner",
            skill_version="report-" + version,
            session_id=run_id,
            skill_path=Path(run["frozen_skill_path"]),
            model=model or self.settings.model,
            parallel=1,
            timeout_seconds=self.settings.timeout_seconds,
            stall_timeout_seconds=self.settings.stall_timeout_seconds,
            max_report_retries=self.settings.max_report_retries,
            output_contract=ReportOutputContract(min_bytes=self.settings.min_report_bytes),
            command=self.settings.command,
            workbuddy_home=self.settings.workbuddy_home,
            product_config=self.settings.product_config,
            allowed_material_roots=(Path(run["data_path"]).parent,),
            openharness_case_ids=(run["case_id"],),
            persist_report_text=False,
        )
        batch = self.runner_func(request)
        payload = (
            batch.to_dict(include_report_text=False)
            if hasattr(batch, "to_dict")
            else batch
        )
        case_result = next(
            (
                item for item in payload.get("cases", [])
                if str(item.get("openharness_case_id")) == run["case_id"]
            ),
            None,
        )
        report_text = report_text_from_result(
            case_result or {},
            payload.get("output_dir") or request.output_root,
        ).strip()
        if not case_result or case_result.get("status") != "generated" or not report_text:
            raise ReportLoopError("Runner 未生成有效报告: %s" % payload.get("status"))
        trace_output_dir = str(payload.get("output_dir") or "")
        try:
            trace_output_dir = (
                Path(trace_output_dir).resolve().relative_to(run_directory).as_posix()
            )
        except (OSError, ValueError):
            pass
        trace = {
            "output_dir": trace_output_dir,
            "case": self._compact_runner_trace(case_result, run_directory),
        }
        return report_text, trace, payload.get("generation_id")

    def _generate_via_api(self, run, version, parent_version, case_file, backend, model, reasoning_effort=None):
        prompt = self._build_report_generation_prompt(run, case_file)
        try:
            response = self.call_llm_func(
                prompt, backend=backend, model=model, reasoning_effort=reasoning_effort
            )
        except Exception as exc:  # noqa: BLE001
            raise ReportLoopError("bianxie API 调用失败: %s" % exc)
        if not isinstance(response, str) or not response.strip():
            raise ReportLoopError("bianxie API 返回了空报告")
        trace = {
            "output_dir": "",
            "case": {
                "status": "generated",
                "backend": backend,
                "model": model,
                "method": "bianxie-api",
                "reasoning_effort": reasoning_effort,
            },
        }
        return response.strip(), trace, None

    def _build_report_generation_prompt(self, run, case_file):
        skill_md = ""
        instructions_md = ""
        skill_root = Path(run["frozen_skill_path"])
        skill_path = skill_root / "SKILL.md"
        instructions_path = skill_root / "references" / "instructions.md"
        if skill_path.is_file():
            skill_md = skill_path.read_text(encoding="utf-8-sig", errors="ignore")
        if instructions_path.is_file():
            instructions_md = instructions_path.read_text(
                encoding="utf-8-sig", errors="ignore"
            )
        case = self._load_case_from_file(run, case_file)
        materials = self._read_materials(case, run)
        turns = case.get("turns") or []
        parts = [
            "你是一名资深行业研究员。请依据下方的「技能规范」「参考资料」与「撰写任务」，"
            "产出一份完整、可直接交付的调研报告，以 Markdown 呈现。\n"
            "仅输出报告正文，不要输出任何修改说明、思考过程或元信息。",
        ]
        if skill_md.strip():
            parts.append("=== 技能规范（Skill）===\n" + skill_md)
        if instructions_md.strip():
            parts.append(
                "=== 核心执行指令（references/instructions.md）===\n"
                + instructions_md
            )
        if materials:
            parts.append("=== 参考资料 ===")
            for name, content in materials:
                parts.append("# 资料：%s\n%s" % (name, content))
        if turns:
            parts.append("=== 撰写任务 ===")
            for idx, turn in enumerate(turns):
                label = turn.get("label") or "turn"
                turn_prompt = turn.get("prompt") or ""
                parts.append("## 第 %d 轮（%s）\n%s" % (idx + 1, label, turn_prompt))
        parts.append("请直接输出完整报告（Markdown）。")
        return "\n\n".join(parts)

    def _load_case_from_file(self, run, case_file):
        path = Path(case_file)
        text = path.read_text(encoding="utf-8-sig", errors="ignore")
        data = json.loads(text)
        if isinstance(data, dict) and isinstance(data.get("cases"), list):
            for case in data["cases"]:
                if str(case.get("case_id")) == str(run["case_id"]):
                    return case
            if data["cases"]:
                return data["cases"][0]
        for line in text.splitlines():
            line = line.strip()
            if not line:
                continue
            try:
                case = json.loads(line)
            except json.JSONDecodeError:
                continue
            if str(case.get("case_id")) == str(run["case_id"]):
                return case
        raise ReportLoopError("无法从 case 文件解析 Report Case: %s" % case_file)

    # 文本类素材：直接按字符读取（兼容 UTF-8 / GBK 等含 BOM 文本）
    _TEXT_SUFFIXES = {".json", ".md", ".csv", ".txt", ".text", ".tsv",
                      ".yml", ".yaml", ".py", ".html", ".xml"}

    def _read_materials(self, case, run):
        """按文件类型真正地把素材转成纯文本，而不是对二进制做裸 read_text。

        - json/md/csv/txt 等：直接读取文本
        - xlsx/xls：用 openpyxl 展开为「工作表 + 行」的表格文本
        - pdf：用 pdfplumber 逐页抽取文本
        - docx：用 python-docx 抽取段落与表格
        - 其它二进制（无可用解析器）：跳过，避免把 ZIP/二进制垃圾喂给模型
        """
        package_root = Path(run["data_path"]).parent.resolve()
        materials = []
        for spec in case.get("input_files") or []:
            if not isinstance(spec, dict) or not spec.get("source"):
                continue
            source = Path(str(spec["source"]))
            if not source.is_absolute():
                source = (package_root / source).resolve()
            if not source.exists():
                continue
            name = spec.get("name") or source.name
            try:
                content = self._read_material_text(source)
            except OSError:
                continue
            if not content or not content.strip():
                continue
            if len(content) > 200000:
                content = content[:200000] + "\n... [资料过长已截断]"
            materials.append((name, content))
        return materials

    def _read_material_text(self, path: Path):
        suffix = path.suffix.lower()
        if suffix in self._TEXT_SUFFIXES:
            return path.read_text(encoding="utf-8-sig", errors="ignore")
        if suffix in (".xlsx", ".xls"):
            return self._xlsx_to_text(path)
        if suffix == ".pdf":
            return self._pdf_to_text(path)
        if suffix == ".docx":
            return self._docx_to_text(path)
        # 未知/无解析器的二进制：跳过，不喂垃圾给模型
        return None

    def _xlsx_to_text(self, path: Path):
        from openpyxl import load_workbook

        try:
            wb = load_workbook(path, data_only=True, read_only=True)
        except Exception:
            return None
        sheets = []
        for ws in wb.worksheets:
            lines = []
            for row in ws.iter_rows(values_only=True):
                cells = ["" if v is None else str(v) for v in row]
                if any(c.strip() for c in cells):
                    lines.append(" | ".join(cells))
            if lines:
                sheets.append(
                    "# 工作表：%s（共 %d 行有效数据）\n%s"
                    % (ws.title, len(lines), "\n".join(lines))
                )
        return "\n\n".join(sheets)

    def _pdf_to_text(self, path: Path):
        import pdfplumber

        chunks = []
        with pdfplumber.open(path) as pdf:
            for idx, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                if text.strip():
                    chunks.append("# 第 %d 页\n%s" % (idx, text))
        return "\n\n".join(chunks)

    def _docx_to_text(self, path: Path):
        from docx import Document

        doc = Document(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    def _judge(
        self,
        run_id: str,
        job_id: str,
        version: str,
        backend: Optional[str],
        model: Optional[str],
        reasoning_effort: Optional[str],
    ) -> None:
        with self._lock:
            run = copy.deepcopy(self._get(run_id))
        item = revision(run, version)
        report_text = self.store.load_report(run_id, version)
        if sha256_text(report_text) != item["report_sha256"]:
            raise ReportLoopError("待评测报告内容与登记哈希不一致")
        case = copy.deepcopy(run["case"])
        preview = run.get("dataset_preview") or {}
        case["structured_data"] = preview.get("structured_data")
        result = self.judge_func(
            case,
            report_text,
            run["rubric"],
            build_judge_prompt,
            lambda prompt: self.call_llm_func(
                prompt,
                backend=backend,
                model=model,
                reasoning_effort=reasoning_effort,
            ),
            llm_client.extract_json,
            strategy=JUDGE_STRATEGY_PER_DIMENSION,
            max_retries=DEFAULT_JUDGE_MAX_RETRIES,
        )

        if result.get("status") != "judged":
            raise ReportLoopError("Judge 未完成: %s" % (result.get("error") or result.get("status")))
        scores = score_labeled_check_judgment(result.get("checks") or {}, run["rubric"])
        check_scores = normalize_check_scores(result.get("checks") or {})
        judgment = {
            "version": version,
            "report_sha256": item["report_sha256"],
            "rubric_sha256": run["rubric_sha256"],
            "checks": result.get("checks") or {},
            "check_scores": check_scores,
            "reasoning": result.get("reasoning") or {},
            "dimensions": scores.get("scores") or {},
            "overall": scores.get("overall"),
            "redline_checks": scores.get("redline_checks") or [],
            "hard_floor_failures": scores.get("hard_floor_failures") or [],
            "case_failed_gate": scores.get("case_failed_gate", False),
            "score_source": "app/report_scoring.py",
            "check_score_source": "app/report_scoring.py:normalize_check_scores",
            "judge_meta": result.get("judge_meta") or {},
            "judge_trace": result.get("judge_trace") or {},
            "backend": backend,
            "model": model,
            "reasoning_effort": reasoning_effort,
            "created_at": _now(),
        }
        failure_report = failure_report_from_checks(
            judgment["checks"], judgment["reasoning"], run["rubric"]
        )
        self.store.save_judgment(run_id, version, judgment)
        with self._lock:
            current = self._get(run_id)
            candidate = revision(current, version)
            if candidate.get("status") != "report_ready":
                raise ReportLoopError("报告版本已被评测或状态已变化")
            candidate["judgment"] = judgment
            candidate["failure_report"] = failure_report
            settlement = settle_judged_revision(current, candidate)
            self.store.append_event(run_id, "report.judged", {
                "version": version,
                "overall": judgment["overall"],
                **settlement,
            })
            self.store.save(current)




